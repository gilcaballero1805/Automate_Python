import docx
import openpyxl
from docx2pdf import convert
# Agrega esta línea al comienzo de tu código
campos_a_reemplazar = ["<Tipo de procedimiento>", "<Num de procedimiento>", "<Objeto>"]

def reemplazar_campos_encabezado(documento, datos):
    seccion = documento.sections[0]
    encabezado = seccion.header
    tabla = encabezado.tables[0]

    # Formato para los dos primeros campos
    formato_dos_primeros = {
        'font_name': 'Arial Nova',
        'font_size': 8,
        'bold': True,
        'underline': True
    }

    # Formato para el tercer campo
    formato_tercer_campo = {
        'font_name': 'Arial Nova',
        'font_size': 8,
        'bold': True,
        'underline': False
    }

    celda = tabla.cell(0, 2)

    # Asegurarse de que haya suficientes párrafos en la celda
    while len(celda.paragraphs) < 3:
        celda.add_paragraph()

    for i, campo in enumerate(["<Tipo de procedimiento>", "<Num de procedimiento>"]):
        parrafo = celda.paragraphs[i]
        parrafo.clear()
        run = parrafo.add_run(datos[campo])

        run.font.name = formato_dos_primeros['font_name']
        run.font.size = docx.shared.Pt(formato_dos_primeros['font_size'])
        run.bold = formato_dos_primeros['bold']
        run.underline = formato_dos_primeros['underline']
        parrafo.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    campo = "<Objeto>"
    parrafo = celda.paragraphs[2]
    parrafo.clear()
    run = parrafo.add_run(datos[campo])

    run.font.name = formato_tercer_campo['font_name']
    run.font.size = docx.shared.Pt(formato_tercer_campo['font_size'])
    run.bold = formato_tercer_campo['bold']
    run.underline = formato_tercer_campo['underline']
    parrafo.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER


def reemplazar_texto_en_parrafo(paragraph, campo_buscar, texto_nuevo, fuente='Arial Nova', tamano=10, bold=False):
    if campo_buscar in paragraph.text:
        # Lista para almacenar los runs del párrafo
        runs = []

        # Obtener los estilos actuales de los runs del párrafo
        for run in paragraph.runs:
            current_style = {
                'font': run.font.name,
                'size': run.font.size.pt,
                'bold': run.bold
            }
            runs.append((run.text, current_style))

        # Reemplazar el texto antiguo por el nuevo
        text = paragraph.text.replace(campo_buscar, texto_nuevo)
        paragraph.clear()

        # Agregar los nuevos runs al párrafo
        for run_text, current_style in runs:
            new_run = paragraph.add_run(run_text)
            new_run.bold = current_style['bold']
            new_run.font.name = current_style['font']
            new_run.font.size = docx.shared.Pt(current_style['size'])

        # Aplicar el formato deseado al nuevo texto
        for run in paragraph.runs:
            run.font.name = fuente
            run.font.size = docx.shared.Pt(tamano)
            if bold:
                run.font.bold = True

def reemplazar_campos_en_documento(documento, campo_buscar, texto_nuevo, bold=False):
    for p in documento.paragraphs:
        for run in p.runs:
            run.font.name = "Arial Nova"
            run.font.size = docx.shared.Pt(10)
            if "Con relación al procedimiento de contratación número No." in run.text:
                run.font.bold = False
        if campo_buscar in p.text:
            # Crear una lista de Runs para el párrafo
            runs = []

            # Dividir el párrafo por el campo_buscar
            partes = p.text.split(campo_buscar)

            for i, parte in enumerate(partes):
                if i > 0:
                    run = p.add_run(texto_nuevo)
                    if bold:
                        run.bold = True
                    runs.append(run)

                run = p.add_run(parte)
                runs.append(run)

            # Eliminar el párrafo original
            p.clear()

            # Agregar los nuevos Runs al párrafo
            for run in runs:
                new_run = p.add_run(run.text)
                new_run.bold = run.bold
                new_run.font.name = run.font.name
                if run.font.size:
                    new_run.font.size = docx.shared.Pt(int(run.font.size.pt))

            # Eliminar el campo_buscar en el último Run agregado
            p.runs[-1].text = p.runs[-1].text.replace(campo_buscar, '')


def leer_datos_excel(ruta_archivo_excel):
    wb = openpyxl.load_workbook(ruta_archivo_excel)
    hoja = wb.active

    datos = {}
    for i, fila in enumerate(hoja.iter_rows(min_row=1, max_row=3, min_col=2, max_col=2, values_only=True)):
        campo_buscar = campos_a_reemplazar[i]
        texto_nuevo = str(fila[0]).rstrip()  # Convertir el valor a texto y eliminar espacios y caracteres de nueva línea al final
        datos[campo_buscar] = texto_nuevo

    return datos

ruta_archivo = "D:\\D2_Python\\2.2_CAPFIN-2_Carta_de_Línea_de_Crédito_b.docx"
documento = docx.Document(ruta_archivo)

ruta_archivo_excel = "D:\\D2_Python\\Anexos.xlsx"
datos_a_reemplazar = leer_datos_excel(ruta_archivo_excel)

reemplazar_campos_encabezado(documento, datos_a_reemplazar)

for campo_buscar, texto_nuevo in datos_a_reemplazar.items():
    reemplazar_campos_en_documento(documento, campo_buscar, texto_nuevo, bold=True)


def cambiar_estilo_primer_parrafo(documento, fuente='Arial Nova', tamano=10):
    primer_parrafo = documento.paragraphs[0]
    primer_parrafo.style.font.name = fuente
    primer_parrafo.style.font.size = docx.shared.Pt(tamano)

# ... (el resto del código original) ...

# Llamar a la función cambiar_estilo_primer_parrafo después de reemplazar campos en el documento
cambiar_estilo_primer_parrafo(documento)

# Cambiar la ruta del archivo de salida a PDF
ruta_archivo_salida = "D:\\D2_Python\\2.2_CAPFIN-2_Carta_de_Línea_de_Crédito_modificado.pdf"

# Guardar el documento modificado como docx
documento.save("D:\\D2_Python\\2.2_CAPFIN-2_Carta_de_Línea_de_Crédito_modificado.docx")

# Convertir el documento docx a PDF
convert("D:\\D2_Python\\2.2_CAPFIN-2_Carta_de_Línea_de_Crédito_modificado.docx", ruta_archivo_salida)


